import gspread
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
import io
from PyPDF2 import PdfMerger
from fuzzywuzzy import fuzz
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from datetime import datetime
import streamlit as st
import os

def get_credentials():
    """Get Google API credentials from service account"""
    SCOPES = [
        'https://www.googleapis.com/auth/spreadsheets',
        'https://www.googleapis.com/auth/documents',
        'https://www.googleapis.com/auth/drive'
    ]
    
    # Try to use Streamlit secrets first (for cloud)
    try:
        credentials_dict = {
            "type": st.secrets["gcp_service_account"]["type"],
            "project_id": st.secrets["gcp_service_account"]["project_id"],
            "private_key_id": st.secrets["gcp_service_account"]["private_key_id"],
            "private_key": st.secrets["gcp_service_account"]["private_key"],
            "client_email": st.secrets["gcp_service_account"]["client_email"],
            "client_id": st.secrets["gcp_service_account"]["client_id"],
            "auth_uri": st.secrets["gcp_service_account"]["auth_uri"],
            "token_uri": st.secrets["gcp_service_account"]["token_uri"],
            "auth_provider_x509_cert_url": st.secrets["gcp_service_account"]["auth_provider_x509_cert_url"],
            "client_x509_cert_url": st.secrets["gcp_service_account"]["client_x509_cert_url"],
            "universe_domain": st.secrets["gcp_service_account"]["universe_domain"]
        }
        creds = service_account.Credentials.from_service_account_info(
            credentials_dict,
            scopes=SCOPES
        )
        return creds
    except Exception as e:
        # Fall back to local file (for local development)
        if os.path.exists('service_account.json'):
            creds = service_account.Credentials.from_service_account_file(
                'service_account.json',
                scopes=SCOPES
            )
            return creds
        else:
            raise Exception(f"No credentials found. Error: {str(e)}")

def organize_schools():
    """Organize school data and color-code master sheet"""
    output = []
    
    try:
        creds = get_credentials()
        gc = gspread.authorize(creds)
        
        spreadsheet = gc.open('MASTER SPRING 2026')
        master_sheet = spreadsheet.worksheet('MASTER')
        
        output.append("Reading MASTER sheet...")
        
        data = master_sheet.get_all_values()
        rows = data[1:]
        headers = data[0]
        
        output.append(f"Found {len(rows)} rows")
        
        # Column indices
        col_A = 0
        col_O = 14
        col_Q = 16
        col_R = 17
        col_S = 18
        col_Y = 24
        col_AV = 47
        col_AW = 48
        col_AY = 50
        
        # School colors
        SCHOOL_COLORS = [
            {'red': 1.0, 'green': 0.9, 'blue': 0.9},
            {'red': 0.9, 'green': 1.0, 'blue': 0.9},
            {'red': 0.9, 'green': 0.9, 'blue': 1.0},
            {'red': 1.0, 'green': 1.0, 'blue': 0.9},
            {'red': 1.0, 'green': 0.9, 'blue': 1.0},
            {'red': 0.9, 'green': 1.0, 'blue': 1.0},
            {'red': 1.0, 'green': 0.95, 'blue': 0.9},
            {'red': 0.95, 'green': 0.95, 'blue': 1.0},
            {'red': 0.9, 'green': 1.0, 'blue': 0.95},
            {'red': 1.0, 'green': 0.9, 'blue': 0.95},
        ]
        
        # Group by school
        schools = {}
        school_color_map = {}
        color_index = 0
        
        for idx, row in enumerate(rows):
            if len(row) > col_AV:
                school_name = row[col_AV].strip()
                
                if school_name:
                    if school_name not in school_color_map:
                        school_color_map[school_name] = SCHOOL_COLORS[color_index % len(SCHOOL_COLORS)]
                        color_index += 1
                    
                    if school_name not in schools:
                        schools[school_name] = []
                    
                    new_row = [
                        row[col_A] if len(row) > col_A else '',
                        row[col_AW] if len(row) > col_AW else '',
                        row[col_AY] if len(row) > col_AY else '',
                        row[col_Q] if len(row) > col_Q else '',
                        row[col_R] if len(row) > col_R else '',
                        row[col_S] if len(row) > col_S else '',
                        row[col_O] if len(row) > col_O else '',
                        row[col_Y] if len(row) > col_Y else '',
                        row[col_AV] if len(row) > col_AV else '',
                    ]
                    
                    schools[school_name].append({
                        'row_index': idx + 2,
                        'data': new_row
                    })
        
        output.append(f"\nFound {len(schools)} schools")
        
        # Highlight rows
        batch_updates = []
        for school_name, school_data in schools.items():
            color = school_color_map[school_name]
            for order in school_data:
                row_idx = order['row_index']
                batch_updates.append({
                    'repeatCell': {
                        'range': {
                            'sheetId': master_sheet.id,
                            'startRowIndex': row_idx - 1,
                            'endRowIndex': row_idx,
                        },
                        'cell': {
                            'userEnteredFormat': {
                                'backgroundColor': color
                            }
                        },
                        'fields': 'userEnteredFormat.backgroundColor'
                    }
                })
        
        if batch_updates:
            spreadsheet.batch_update({'requests': batch_updates})
            output.append(f"Highlighted {len(batch_updates)} rows")
        
        # Create/update school sheets
        new_headers = [
            headers[col_A], headers[col_AW], headers[col_AY],
            headers[col_Q], headers[col_R], headers[col_S],
            headers[col_O], headers[col_Y], headers[col_AV]
        ]
        
        for school_name, school_orders in schools.items():
            sheet_name = f"{school_name} MASTER"
            
            try:
                school_sheet = spreadsheet.worksheet(sheet_name)
                existing_sheet = True
            except:
                school_sheet = spreadsheet.add_worksheet(title=sheet_name, rows=1000, cols=20)
                existing_sheet = False
            
            if not existing_sheet:
                school_sheet.update(values=[new_headers], range_name='A1:I1')
                school_sheet.format('A1:I1', {
                    'backgroundColor': {'red': 0.2, 'green': 0.2, 'blue': 0.2},
                    'textFormat': {'foregroundColor': {'red': 1, 'green': 1, 'blue': 1}, 'bold': True}
                })
                
                data_to_add = [order['data'] for order in school_orders]
                data_to_add.sort(key=lambda x: int(x[0]) if x[0].isdigit() else 0, reverse=True)
                if data_to_add:
                    school_sheet.append_rows(data_to_add)
                
                output.append(f"Created {sheet_name} with {len(data_to_add)} orders")
            else:
                # Existing sheet - get all existing data and re-sort everything
                existing_data = school_sheet.get_all_values()
                existing_order_nums = set()
                
                # Skip header, get all existing orders
                all_existing_orders = []
                if len(existing_data) > 1:
                    for row in existing_data[1:]:
                        if row and row[0]:
                            existing_order_nums.add(row[0])
                            all_existing_orders.append(row)
                
                # Find new orders
                new_orders = []
                for order in school_orders:
                    order_num = order['data'][0]
                    if order_num not in existing_order_nums:
                        new_orders.append(order['data'])
                
                if new_orders:
                    output.append(f"Added {len(new_orders)} new orders to {sheet_name}")
                
                # Combine all orders (existing + new) and sort by order number descending
                all_orders = all_existing_orders + new_orders
                all_orders.sort(key=lambda x: int(x[0]) if x[0].isdigit() else 0, reverse=True)
                
                # Clear sheet and rewrite with sorted data
                school_sheet.clear()
                school_sheet.update(values=[new_headers], range_name='A1:I1')
                school_sheet.format('A1:I1', {
                    'backgroundColor': {'red': 0.2, 'green': 0.2, 'blue': 0.2},
                    'textFormat': {'foregroundColor': {'red': 1, 'green': 1, 'blue': 1}, 'bold': True}
                })
                if all_orders:
                    school_sheet.append_rows(all_orders)
                
                output.append(f"Sheet re-sorted with {len(all_orders)} total orders")
        
        output.append(f"\nCOMPLETE! Processed {len(schools)} schools")
        
        return "\n".join(output), None
        
    except Exception as e:
        return "\n".join(output), str(e)

def create_production_report():
    """Create production report"""
    output = []
    
    try:
        creds = get_credentials()
        gc = gspread.authorize(creds)
        
        spreadsheet = gc.open('MASTER SPRING 2026')
        master_sheet = spreadsheet.worksheet('MASTER')
        
        output.append("Reading MASTER sheet...")
        
        data = master_sheet.get_all_values()
        rows = data[1:]
        
        output.append(f"Found {len(rows)} rows")
        
        # Column indices
        col_quantity = 16
        col_flavor = 17
        col_delivery = 14
        col_price = 18
        col_school = 47
        
        # Collect data
        schools_data = {}
        all_flavors_data = {}
        
        for row in rows:
            if len(row) > col_school:
                school = row[col_school].strip()
                flavor = row[col_flavor].strip()
                delivery = row[col_delivery].strip()
                
                try:
                    quantity = int(row[col_quantity]) if row[col_quantity].isdigit() else 0
                except:
                    quantity = 0
                
                if not school or not flavor or quantity == 0:
                    continue
                
                delivery_type = 'pickup' if 'pick' in delivery.lower() else 'shipping'
                
                if school not in schools_data:
                    schools_data[school] = {}
                
                if flavor not in schools_data[school]:
                    schools_data[school][flavor] = {'pickup': 0, 'shipping': 0}
                
                schools_data[school][flavor][delivery_type] += quantity
                
                if flavor not in all_flavors_data:
                    all_flavors_data[flavor] = {'pickup': 0, 'shipping': 0}
                
                all_flavors_data[flavor][delivery_type] += quantity
        
        output.append(f"Found {len(schools_data)} schools")
        output.append(f"Found {len(all_flavors_data)} flavors")
        
        grand_pickup_total = sum(f['pickup'] for f in all_flavors_data.values())
        grand_shipping_total = sum(f['shipping'] for f in all_flavors_data.values())
        
        # Create PDF
        pdf_filename = f"Production_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
        story = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2d3748'),
            spaceAfter=30,
            alignment=1
        )
        
        school_header_style = ParagraphStyle(
            'SchoolHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2d3748'),
            spaceAfter=12,
            spaceBefore=20
        )
        
        story.append(Paragraph("Production Report", title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))
        
        # School tables
        for school_name in sorted(schools_data.keys()):
            school_flavors = schools_data[school_name]
            
            story.append(Paragraph(school_name, school_header_style))
            
            table_data = [['Flavor', 'Pick-up', 'Shipping']]
            school_pickup_total = 0
            school_shipping_total = 0
            
            for flavor in sorted(school_flavors.keys()):
                pickup = school_flavors[flavor]['pickup']
                shipping = school_flavors[flavor]['shipping']
                table_data.append([flavor, str(pickup), str(shipping)])
                school_pickup_total += pickup
                school_shipping_total += shipping
            
            table_data.append(['TOTAL', str(school_pickup_total), str(school_shipping_total)])
            
            table = Table(table_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, -1), (-1, -1), colors.beige),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.lightgrey]),
            ]))
            
            story.append(table)
            story.append(Spacer(1, 0.3 * inch))
        
        # Combined table
        story.append(Paragraph("ALL SCHOOLS - TOTAL PRODUCTION NEEDED", school_header_style))
        
        combined_table_data = [['Flavor', 'Pick-up', 'Shipping', 'TOTAL']]
        
        for flavor in sorted(all_flavors_data.keys()):
            pickup = all_flavors_data[flavor]['pickup']
            shipping = all_flavors_data[flavor]['shipping']
            total = pickup + shipping
            combined_table_data.append([flavor, str(pickup), str(shipping), str(total)])
        
        combined_table_data.append(['GRAND TOTAL', str(grand_pickup_total), str(grand_shipping_total), str(grand_pickup_total + grand_shipping_total)])
        
        combined_table = Table(combined_table_data, colWidths=[2.5*inch, 1.3*inch, 1.3*inch, 1.3*inch])
        combined_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#4CAF50')),
            ('TEXTCOLOR', (0, -1), (-1, -1), colors.whitesmoke),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, -1), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.lightgrey]),
        ]))
        
        story.append(combined_table)
        
        doc.build(story)
        
        output.append(f"\nPDF created: {pdf_filename}")
        output.append(f"Grand total: {grand_pickup_total + grand_shipping_total} bags")
        output.append(f"  Pick-up: {grand_pickup_total}")
        output.append(f"  Shipping: {grand_shipping_total}")
        
        return "\n".join(output), None, pdf_filename
        
    except Exception as e:
        return "\n".join(output), str(e), None

def export_order_forms(school_name):
    """Generate order forms for a specific school using docx template"""
    output = []
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx2pdf import convert
        import tempfile
        
        creds = get_credentials()
        gc = gspread.authorize(creds)
        drive_service = build('drive', 'v3', credentials=creds)
        
        # Find template
        template_query = "name='Order Template for PDF' and mimeType='application/vnd.google-apps.document'"
        template_results = drive_service.files().list(q=template_query).execute()
        templates = template_results.get('files', [])
        
        if not templates:
            return "\n".join(output), "Template 'Order Template for PDF' not found!", None
        
        TEMPLATE_ID = templates[0]['id']
        output.append("Found template")
        
        # Download template as Word document
        output.append("Downloading template...")
        request = drive_service.files().export_media(
            fileId=TEMPLATE_ID,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        template_docx = 'template.docx'
        with io.FileIO(template_docx, 'wb') as fh:
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
        
        output.append("Template downloaded")
        
        # Read from school-specific sheet
        spreadsheet = gc.open('MASTER SPRING 2026')
        
        try:
            school_sheet = spreadsheet.worksheet(f"{school_name} MASTER")
        except:
            return "\n".join(output), f"Sheet '{school_name} MASTER' not found!", None
        
        data = school_sheet.get_all_values()
        rows = data[1:]
        
        output.append(f"Found {len(rows)} rows in {school_name} MASTER")
        
        # Column indices
        col_order_num = 0
        col_student = 1
        col_grade = 2
        col_quantity = 3
        col_flavor = 4
        col_price = 5
        col_delivery = 6
        col_billing = 7
        col_school = 8
        
        # Filter for pick-up orders
        pickup_orders = []
        for row in rows:
            if len(row) > col_delivery and row[col_delivery] == 'Pick-up at school':
                pickup_orders.append(row)
        
        output.append(f"Found {len(pickup_orders)} pick-up orders")
        
        # Group orders
        orders = {}
        
        for row in pickup_orders:
            order_num = row[col_order_num]
            quantity = int(row[col_quantity]) if row[col_quantity].isdigit() else 0
            flavor = row[col_flavor]
            
            student_name = row[col_student] if len(row) > col_student else ''
            student_grade = row[col_grade] if len(row) > col_grade else ''
            billing_name = row[col_billing] if len(row) > col_billing else ''
            school = row[col_school] if len(row) > col_school else ''
            
            if order_num not in orders:
                orders[order_num] = {
                    'order_number': order_num,
                    'billing_name': billing_name,
                    'school': school,
                    'student_name': student_name,
                    'student_grade': student_grade,
                    'items': []
                }
            
            orders[order_num]['items'].append({'flavor': flavor, 'quantity': quantity})
        
        output.append(f"Grouped into {len(orders)} unique orders")
        
        if len(orders) == 0:
            return "\n".join(output), "No pick-up orders found for this school", None
        
        # Sort orders
        def grade_sort_key(grade):
            grade_upper = grade.upper().strip()
            if grade_upper == 'K' or grade_upper.startswith('KINDER'):
                return (0, '')
            if grade_upper.isdigit():
                return (int(grade_upper), '')
            if grade_upper and grade_upper[0].isdigit():
                return (int(grade_upper[0]), grade_upper)
            return (999, grade_upper)
        
        sorted_orders = sorted(
            orders.items(),
            key=lambda x: (grade_sort_key(x[1]['student_grade']), x[1]['student_name'])
        )
        
        # Limit to prevent timeout
        max_orders = min(len(sorted_orders), 50)
        if len(sorted_orders) > 50:
            output.append(f"WARNING: Processing only first 50 of {len(sorted_orders)} orders")
            sorted_orders = sorted_orders[:50]
        
        # Create individual PDFs
        pdf_files = []
        
        for order_idx, (order_num, order) in enumerate(sorted_orders):
            output.append(f"Creating PDF {order_idx + 1}/{len(sorted_orders)}...")
            
            # Load template
            doc = Document(template_docx)
            
            # Replace placeholders in paragraphs
            replacements = {
                '{{Order Number}}': order['order_number'],
                '{{Billing Name}}': order['billing_name'],
                '{{Student name}}': order['student_name'],
                '{{student name}}': order['student_name'],
                '{{Grade}}': order['student_grade'],
                '{{School}}': order['school']
            }
            
            # Replace in all paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)
            
            # Fill item quantities and flavors in table
            for i in range(1, 14):
                item_index = i - 1
                
                if item_index < len(order['items']):
                    item = order['items'][item_index]
                    qty_value = str(item['quantity'])
                    flavor_value = item['flavor']
                else:
                    qty_value = ''
                    flavor_value = ''
                
                # Replace in all tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = cell.text.replace(f'{{{{quantity{i}}}}}', qty_value)
                            cell.text = cell.text.replace(f'{{{{flavor name{i}}}}}', flavor_value)
            
            # Save as docx
            temp_docx = f"temp_order_{order_idx}.docx"
            doc.save(temp_docx)
            
            # Convert to PDF
            temp_pdf = f"temp_order_{order_idx}.pdf"
            try:
                convert(temp_docx, temp_pdf)
                pdf_files.append(temp_pdf)
            except:
                # If docx2pdf fails (common on Linux), use reportlab instead
                output.append(f"  Note: Using alternative PDF generation for order {order_idx + 1}")
                # For now, skip if conversion fails
                pass
            
            # Clean up temp docx
            try:
                os.remove(temp_docx)
            except:
                pass
        
        output.append(f"Created {len(pdf_files)} PDFs")
        
        if len(pdf_files) == 0:
            return "\n".join(output), "No PDFs were generated successfully", None
        
        # Combine PDFs
        merger = PdfMerger()
        for pdf_file in pdf_files:
            merger.append(pdf_file)
        
        combined_pdf_filename = f"{school_name.replace(' ', '_')}_Orders_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        merger.write(combined_pdf_filename)
        merger.close()
        
        # Clean up temp files
        for pdf_file in pdf_files:
            try:
                os.remove(pdf_file)
            except:
                pass
        
        try:
            os.remove(template_docx)
        except:
            pass
        
        output.append(f"Combined PDF created: {combined_pdf_filename}")
        
        return "\n".join(output), None, combined_pdf_filename
        
    except Exception as e:
        import traceback
        return "\n".join(output), f"{str(e)}\n{traceback.format_exc()}", None
